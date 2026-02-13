"""
Document processing service.

Primary extraction via Docling (IBM) — AI-powered layout analysis, table structure
(TableFormer), and built-in OCR. Outputs Markdown with headings, tables, bold, lists.
Supports PDF, DOCX, PPTX, HTML, XLSX natively.

Legacy fallback chain (PyMuPDF → pdfplumber → pypdf) retained for environments
where Docling is not installed.

Also supports multilingual document processing including:
- Space-separated languages: English, Swedish, Finnish, Spanish, German, French
- CJK languages: Chinese, Japanese (proper sentence splitting)
"""

import os
import re
import csv
import json
import logging
import threading
from src.utils.multilingual import split_sentences as multilingual_split_sentences

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Docling singleton converter (avoids reloading ~1-2 GB AI models per file)
# ---------------------------------------------------------------------------
_converter = None
_converter_lock = threading.Lock()

# Formats that Docling handles natively
DOCLING_EXTENSIONS = {".pdf", ".docx", ".pptx", ".html", ".htm", ".xlsx"}


def _get_converter():
    """
    Return a singleton DocumentConverter instance.

    Thread-safe via lock. Lazy-imports docling so the module loads
    even when docling is not installed (legacy fallback still works).
    """
    global _converter
    if _converter is not None:
        return _converter

    with _converter_lock:
        # Double-check after acquiring lock
        if _converter is not None:
            return _converter

        from docling.document_converter import DocumentConverter, PdfFormatOption
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions

        pipeline_options = PdfPipelineOptions(
            do_ocr=True,
            do_table_structure=True,
        )
        _converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
        logger.info("[Docling] DocumentConverter initialized (OCR + table structure enabled)")
        return _converter


def extract_with_docling(file_path: str) -> str:
    """
    Extract document content as Markdown using Docling.

    Handles PDF (with OCR + table structure), DOCX, PPTX, HTML, XLSX.
    Returns Markdown string with ``# headings``, ``| table |`` formatting,
    ``**bold**``, bullet lists, etc.

    Raises:
        Exception: If Docling conversion fails.
    """
    converter = _get_converter()
    result = converter.convert(file_path)
    return result.document.export_to_markdown()


# ---------------------------------------------------------------------------
# Main entry point — read_text()
# ---------------------------------------------------------------------------

def read_text(file_path: str, text_max: int = 400000):
    """
    Read text from various file formats.

    Returns list of (page_num, text) tuples. For Docling-processed formats,
    always returns [(0, markdown_text)] since Docling outputs the full document
    as one Markdown string. Downstream callers (ingestion.py) concatenate all
    pages into full_text anyway.

    Extraction strategy:
    - PDF/DOCX/PPTX/HTML/XLSX: Docling (AI layout + tables + OCR), with
      legacy PyMuPDF/pdfplumber/pypdf fallback for PDF if Docling fails.
    - CSV: csv.reader (simple, correct as-is)
    - JSON: json.load with pretty-print (simple, correct as-is)
    - Plaintext/Markdown: direct read (simple, correct as-is)
    """
    ext = os.path.splitext(file_path)[1].lower()

    # --- Docling-supported formats ---
    if ext in DOCLING_EXTENSIONS:
        try:
            markdown_text = extract_with_docling(file_path)
            if markdown_text and markdown_text.strip():
                logger.info(
                    f"[Docling] Extracted {len(markdown_text)} chars from {os.path.basename(file_path)}"
                )
                return [(0, markdown_text[:text_max])]
            logger.warning(
                f"[Docling] Empty output for {file_path}, trying legacy fallback"
            )
        except Exception as e:
            logger.warning(f"[Docling] Failed for {file_path}: {e}, trying legacy fallback")

        # Legacy fallback — only PDF has legacy extractors
        if ext == ".pdf":
            return _legacy_extract_pdf_with_fallback(file_path)
        # For DOCX, try python-docx legacy
        if ext == ".docx":
            return _legacy_extract_docx(file_path, text_max)
        # No legacy fallback for PPTX/HTML/XLSX — return empty
        logger.warning(f"[Docling] No legacy fallback available for {ext}: {file_path}")
        return []

    # --- Simple formats (no Docling needed) ---
    if ext == ".csv":
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            all_rows = [",".join(row) for row in reader]
            return [(0, "\n".join(all_rows)[:text_max])]

    if ext == ".json":
        with open(file_path, "r", encoding="utf-8") as f:
            try:
                data = json.load(f)
                return [(0, json.dumps(data, indent=2)[:text_max])]
            except Exception:
                f.seek(0)
                return [(0, f.read()[:text_max])]

    # Plaintext fallback (.txt, .md, .log, etc.)
    with open(file_path, "r", encoding="utf-8") as f:
        return [(0, f.read()[:text_max])]


# ---------------------------------------------------------------------------
# Legacy DOCX extraction (fallback when Docling unavailable)
# ---------------------------------------------------------------------------

def _legacy_extract_docx(file_path: str, text_max: int = 400000) -> list:
    """Extract DOCX using python-docx. Legacy fallback for when Docling is unavailable."""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []
        for element in doc.element.body:
            if element.tag.endswith("p"):
                for para in doc.paragraphs:
                    if para._element is element and para.text.strip():
                        text_parts.append(para.text)
                        break
            elif element.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._element is element:
                        table_rows = []
                        for row in table.rows:
                            row_text = " | ".join(
                                cell.text.strip() for cell in row.cells
                            )
                            if row_text.strip():
                                table_rows.append(row_text)
                        if table_rows:
                            text_parts.append("[TABLE]\n" + "\n".join(table_rows))
                        break
        return [(0, "\n\n".join(text_parts)[:text_max])]
    except Exception as e:
        logger.warning(f"[Legacy DOCX] Failed for {file_path}: {e}")
        return []


# ---------------------------------------------------------------------------
# Legacy PDF extraction functions (fallback when Docling unavailable)
# ---------------------------------------------------------------------------

def _legacy_clean_pdf_text(text: str) -> str:
    """Clean extracted PDF text by removing artifacts and normalizing whitespace."""
    if not text:
        return ""
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    return text.strip()


def _legacy_extract_with_pymupdf(file_path: str) -> list:
    """Extract text using PyMuPDF (fitz) — legacy fallback."""
    try:
        import fitz

        doc = fitz.open(file_path)
        pages_text = []
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")
            text = _legacy_clean_pdf_text(text)
            if text:
                pages_text.append((page_num, text))
        doc.close()
        if pages_text:
            logger.debug(f"[Legacy PDF] PyMuPDF extracted {len(pages_text)} pages from {file_path}")
        return pages_text
    except Exception as e:
        logger.debug(f"[Legacy PDF] PyMuPDF failed for {file_path}: {e}")
        return []


def _legacy_extract_with_pdfplumber(file_path: str) -> list:
    """Extract text using pdfplumber — legacy fallback for tables/complex layouts."""
    try:
        import pdfplumber

        pages_text = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                tables = page.find_tables()

                if tables:
                    table_bboxes = [table.bbox for table in tables]

                    def not_in_table(obj):
                        obj_x = (obj["x0"] + obj["x1"]) / 2
                        obj_y = (obj["top"] + obj["bottom"]) / 2
                        for bbox in table_bboxes:
                            x0, top, x1, bottom = bbox
                            if x0 <= obj_x <= x1 and top <= obj_y <= bottom:
                                return False
                        return True

                    filtered_page = page.filter(not_in_table)
                    text = filtered_page.extract_text() or ""

                    table_texts = []
                    for table in tables:
                        extracted = table.extract()
                        if extracted:
                            rows = []
                            for row in extracted:
                                if row:
                                    row_text = " | ".join(
                                        str(cell).strip() if cell else ""
                                        for cell in row
                                    )
                                    rows.append(row_text)
                            if rows:
                                table_texts.append("\n".join(rows))

                    if table_texts:
                        text = text + "\n\n[TABLE]\n" + "\n\n[TABLE]\n".join(table_texts)
                else:
                    text = page.extract_text() or ""

                text = _legacy_clean_pdf_text(text)
                if text:
                    pages_text.append((page_num, text))

        if pages_text:
            logger.debug(f"[Legacy PDF] pdfplumber extracted {len(pages_text)} pages from {file_path}")
        return pages_text
    except Exception as e:
        logger.debug(f"[Legacy PDF] pdfplumber failed for {file_path}: {e}")
        return []


def _legacy_extract_with_pypdf(file_path: str) -> list:
    """Extract text using pypdf — legacy fallback."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        pages_text = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = _legacy_clean_pdf_text(text)
            if text:
                pages_text.append((page_num, text))
        if pages_text:
            logger.debug(f"[Legacy PDF] pypdf extracted {len(pages_text)} pages from {file_path}")
        return pages_text
    except Exception as e:
        logger.debug(f"[Legacy PDF] pypdf failed for {file_path}: {e}")
        return []


def _legacy_extract_pdf_with_fallback(file_path: str) -> list:
    """
    Legacy PDF extraction fallback chain:
    1. PyMuPDF (fitz) — fastest, good general extraction
    2. pdfplumber — better for tables and complex layouts
    3. pypdf — last resort
    """
    result = _legacy_extract_with_pymupdf(file_path)
    if result:
        return result

    result = _legacy_extract_with_pdfplumber(file_path)
    if result:
        return result

    result = _legacy_extract_with_pypdf(file_path)
    if result:
        return result

    logger.warning(f"[Legacy PDF] All extraction methods failed for {file_path}")
    return []


# ---------------------------------------------------------------------------
# Sentence splitting & chunking (used by batch_ingest.py)
# ---------------------------------------------------------------------------

def sentence_split(text: str) -> list[str]:
    """
    Split text into sentences with multilingual support.

    Handles:
    - Latin languages: . ! ? followed by space and capital letter
    - Chinese: 。 ！ ？ ；
    - Japanese: 。 ！ ？
    - Mixed multilingual text
    """
    return multilingual_split_sentences(text)


def make_chunks(pages_text: list, target: int = 400, overlap: int = 90) -> list[tuple]:
    """Split document into overlapping chunks"""
    all_chunks = []

    for page_num, text in pages_text:
        chunks, buff, size = [], [], 0
        sentences = sentence_split(text)

        for s in sentences:
            buff.append(s)
            if size + len(s) <= target:
                size += len(s) + 1
            else:
                if buff:
                    chunks.append((page_num, " ".join(buff)))

                overlap_sentences = []
                overlap_size = 0
                for sent in reversed(buff):
                    if (
                        overlap_size + len(sent) + (1 if overlap_sentences else 0)
                        <= overlap
                    ):
                        overlap_sentences.insert(0, sent)
                        overlap_size += len(sent) + (1 if overlap_size > 0 else 0)
                    else:
                        break

                buff = overlap_sentences
                size = sum(len(s) for s in buff) + max(0, len(buff) - 1)

        if buff:
            chunks.append((page_num, " ".join(buff)))

        all_chunks.extend(chunks)

    return all_chunks