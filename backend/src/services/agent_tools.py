"""
Agent tools for the ReAct agent.

This module defines:
1. Tool schemas (what the LLM sees - OpenAI function calling format)
2. Tool execution functions (what actually runs when LLM calls a tool)
3. Tool registry (maps tool names to execution functions)

Each tool has two parts:
- SCHEMA: Describes the tool to the LLM (name, description, parameters)
- execute_xxx: The Python function that actually executes the tool
"""

import os
import re
import inspect
from datetime import datetime
from typing import Dict, Any
import logging

from src.application.services import FileService
from urllib.parse import urlparse
import urllib.parse
import requests
from langsmith import traceable
from src.services.retrieval import build_where
from src.services.retrieval_decomposition import retrieve_with_decomposition
from src.config.settings import Config
from src.utils.safety import scrub_context
from src.services.retrieval_evaluator import RetrievalEvaluator
from src.models.evaluation import ReflectionConfig, EvaluationCriteria
from src.services.query_refiner import QueryRefiner
from src.services.clarification_helper import ClarificationHelper
from src.services.web_search import WebSearchService
from src.services.browser_download import browser_download, is_browser_use_available
from src.utils.send_email import send_email
from src.observability.metrics import increment_error, MetricsErrorType

logger = logging.getLogger(__name__)

# ============================================================================
# TOOL 1: SEARCH DOCUMENTS
# ============================================================================

SEARCH_DOCUMENTS_SCHEMA = {
    "type": "function",
    "function": {
        "name": "search_documents",
        "description": (
            "Search internal company documents for information. "
            "Use this when the user asks about company data, reports, policies, "
            "or any information that might be in uploaded documents. "
            "Returns relevant document chunks with sources and page numbers."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": (
                        "The search query. Be specific and use keywords from the user's question. "
                        "For example: 'Q1 2024 revenue', 'employee vacation policy', 'sales targets'"
                    ),
                },
            },
            "required": ["query"],
        },
    },
}


@traceable
def execute_search_documents(args: Dict[str, Any], context: Dict[str, Any]) -> str:
    """
    Execute the search_documents tool.

    Args:
        args: Arguments from LLM containing:
            - query (str): Search query

        context: System context containing:
            - vector_db: VectorDB instance
            - dept_id: Department ID from auth
            - user_id: User ID from auth
            - request: Flask request object (for filters)

    Returns:
        Formatted string result for LLM containing document chunks

    Flow with refinement loop:
    1. Retrieve contexts with original query
    2. Evaluate quality
    3. If REFINE recommended and attempts < MAX:
       - Refine query
       - Retrieve again
       - Re-evaluate
       - Repeat until ANSWER or max attempts
    4. Return best contexts found
    """
    query = args.get("query", "")
    top_k = Config.TOP_K
    vector_db = context.get("vector_db", None)
    client = context.get("openai_client", None)
    dept_id = context.get("dept_id", "")
    user_id = context.get("user_id", "")
    use_hybrid = context.get("use_hybrid", Config.USE_HYBRID)
    use_reranker = context.get("use_reranker", Config.USE_RERANKER)
    request_data = context.get("request_data", {})
    if not dept_id or not user_id:
        return "Error: No organization ID or user ID provided"
    if not vector_db:
        return "Error: No vector database available"
    if not client:
        return "Error: No OpenAI client available"

    try:
        # Build where clause
        where = build_where(request_data, dept_id, user_id)

        # Initial retrieval
        current_query = query
        ctx, _ = retrieve_with_decomposition(
            vector_db=vector_db,
            openai_client=client,
            query=current_query,
            dept_id=dept_id,
            user_id=user_id,
            top_k=top_k,
            where=where,
            use_hybrid=use_hybrid,
            use_reranker=use_reranker,
        )

        logger.debug(
            f"[SEARCH_DOCUMENTS] First time retrieved {len(ctx)} contexts for query: {current_query}"
        )

        if not ctx:
            return "No relevant documents found."

        # Store raw contexts in context dict for agent to access later
        context["_retrieved_contexts"] = ctx

        # Initialize messages (set by self-reflection based on quality)
        clarification_msg = None
        external_msg = None

        if Config.USE_SELF_REFLECTION:
            config = ReflectionConfig.from_settings(Config)
            evaluator = RetrievalEvaluator(config=config, openai_client=client)

            # Initial evaluation
            criteria = EvaluationCriteria(
                query=current_query,
                contexts=ctx,
                search_metadata={
                    "hybrid": use_hybrid,
                    "reranker": use_reranker,
                    "top_k": top_k,
                },
                mode=config.mode,
            )
            eval_result = evaluator.evaluate(criteria)
            context["_last_evaluation"] = eval_result

            # Log initial evaluation
            print(
                f"[SELF-REFLECTION] Quality: {eval_result.quality.value}, "
                f"Confidence: {eval_result.confidence:.2f}, "
                f"Recommendation: {eval_result.recommendation.value}"
            )
            print(f"[SELF-REFLECTION] Reasoning: {eval_result.reasoning}")
            if eval_result.issues:
                print(f"[SELF-REFLECTION] Issues: {', '.join(eval_result.issues)}")
            if eval_result.missing_aspects:
                print(
                    f"[SELF-REFLECTION] Missing aspects: {', '.join(eval_result.missing_aspects)}"
                )

            # =================================================================
            # Handle Direct CLARIFY (confidence < 0.5) - skip refinement
            # =================================================================
            if eval_result.should_clarify:
                clarifier = ClarificationHelper(openai_client=client)
                clarification_msg = clarifier.generate_clarification(
                    query=query,
                    eval_result=eval_result,
                    max_attempts_reached=False,
                    context_hint="uploaded documents",
                )
                print(f"[CLARIFICATION] Direct clarify (confidence < 0.5)")

            # =================================================================
            # Handle EXTERNAL Recommendation (Week 3 - Day 6)
            # =================================================================
            # Signal agent to use web_search tool (don't execute automatically)
            if eval_result.should_search_external:
                print(f"[SELF-REFLECTION] EXTERNAL recommended - suggesting web search")
                external_msg = (
                    f"[EXTERNAL SEARCH SUGGESTED]\n"
                    f'The query "{query}" appears to require external information '
                    f"not found in the uploaded documents.\n\n"
                    f"Reason: {eval_result.reasoning}\n\n"
                    f"Recommendation: Use the web_search tool to find current or external information.\n\n"
                    f"---\n\n"
                )

            # Refinement loop with local counter (only if not already CLARIFY)
            refinement_count = 0
            if Config.REFLECTION_AUTO_REFINE and eval_result.should_refine:
                max_attempts = Config.REFLECTION_MAX_REFINEMENT_ATTEMPTS
                openai_model = context.get("model", Config.OPENAI_MODEL)
                temperature = context.get("temperature", Config.OPENAI_TEMPERATURE)
                query_refiner = QueryRefiner(
                    openai_client=client, model=openai_model, temperature=temperature
                )

                # Track refinement history for this tool call (for logging)
                refinement_history = []

                while refinement_count < max_attempts:

                    refinement_count += 1
                    print(
                        f"[QUERY_REFINER] Refinement attempt {refinement_count}/{max_attempts}"
                    )

                    # Refine the query (use current_query, not original)
                    refined_query = query_refiner.refine_query(
                        original_query=current_query,
                        eval_result=eval_result,
                        context_hint=" ".join([c.get("chunk", "")[:100] for c in ctx]),
                    )

                    # Track for logging
                    refinement_history.append(
                        {
                            "attempt": refinement_count,
                            "from": current_query,
                            "to": refined_query,
                        }
                    )
                    logger.info(
                        f"[REFINED_QUERY] Original: '{current_query}' → Refined: '{refined_query}'"
                    )

                    # Update current query for next iteration
                    current_query = refined_query

                    # Retrieve with refined query
                    ctx, _ = retrieve_with_decomposition(
                        vector_db=vector_db,
                        openai_client=client,
                        query=current_query,
                        dept_id=dept_id,
                        user_id=user_id,
                        top_k=top_k,
                        where=where,
                        use_hybrid=use_hybrid,
                        use_reranker=use_reranker,
                    )

                    if not ctx:
                        logger.info(
                            "[REFINED_QUERY] No results for refined query, keeping previous results"
                        )
                        # Restore previous contexts
                        ctx = context["_retrieved_contexts"]
                        break

                    # Update stored contexts
                    context["_retrieved_contexts"] = ctx

                    # Re-evaluate
                    criteria = EvaluationCriteria(
                        query=current_query,
                        contexts=ctx,
                        search_metadata={
                            "hybrid": use_hybrid,
                            "reranker": use_reranker,
                            "top_k": top_k,
                        },
                        mode=config.mode,
                    )
                    eval_result = evaluator.evaluate(criteria)
                    context["_last_evaluation"] = eval_result

                    # Log refined evaluation
                    print(
                        f"[SELF-REFLECTION] (Attempt {refinement_count}) Quality: {eval_result.quality.value}, "
                        f"Confidence: {eval_result.confidence:.2f}, "
                        f"Recommendation: {eval_result.recommendation.value}"
                    )

                    # Early exit if quality is now good enough (ANSWER recommendation)
                    if eval_result.should_answer:
                        print(
                            f"[QUERY_REFINER] Quality achieved ANSWER level (confidence: {eval_result.confidence:.2f}), stopping refinement early"
                        )
                        break

                # Log final status
                if refinement_count > 0:
                    if (
                        refinement_count >= max_attempts
                        and not eval_result.should_answer
                    ):
                        # Only show clarification if quality is STILL poor after all attempts
                        print(
                            f"[QUERY_REFINER] Max refinement attempts ({max_attempts}) reached, quality still {eval_result.quality.value}"
                        )

                        # =================================================================
                        # Progressive Fallback: max attempts reached AND still poor → clarification
                        # =================================================================
                        clarifier = ClarificationHelper(openai_client=client)
                        clarification_msg = clarifier.generate_clarification(
                            query=query,  # Original query (not refined)
                            eval_result=eval_result,
                            max_attempts_reached=True,
                            context_hint="uploaded documents",
                        )
                        print(
                            f"[CLARIFICATION] Max attempts reached with poor quality, suggesting clarification"
                        )
                    elif eval_result.should_answer:
                        print(
                            f"[QUERY_REFINER] Refinement successful after {refinement_count} attempt(s), "
                            f"quality: {eval_result.quality.value}, confidence: {eval_result.confidence:.2f}"
                        )
                    else:
                        print(
                            f"[QUERY_REFINER] Refinement ended after {refinement_count} attempt(s), "
                            f"quality: {eval_result.quality.value}"
                        )

                    # Store history in context for debugging (optional)
                    context["_refinement_history"] = refinement_history

        # Continue with normal flow: format and return contexts
        for c in ctx:
            c["chunk"] = scrub_context(c.get("chunk", ""))

        # Format contexts for LLM - group by sub-query if decomposition was used
        # Check if contexts have sub_query labels (indicates decomposition was used)
        # IMPORTANT: Use list with dict.fromkeys() to preserve insertion order from ctx
        # (set doesn't guarantee order, causing citation number mismatch with frontend)
        sub_queries = list(
            dict.fromkeys(c.get("sub_query") for c in ctx if c.get("sub_query"))
        )
        has_decomposition = len(sub_queries) > 1

        if has_decomposition:
            # Group contexts by sub-query for clearer presentation to LLM
            context_str = f'Original Query: "{query}"\n'
            context_str += f"Decomposed into {len(sub_queries)} sub-queries for better retrieval:\n\n"

            context_idx = 1
            for sq in sub_queries:
                sq_contexts = [c for c in ctx if c.get("sub_query") == sq]
                context_str += (
                    f'=== Sub-query: "{sq}" ({len(sq_contexts)} results) ===\n\n'
                )

                for c in sq_contexts:
                    context_str += (
                        f"Context {context_idx} (Source: {c.get('source', 'unknown')}"
                        + (f", Page: {c['page']}" if c.get("page", 0) > 0 else "")
                        + (f", file_id: {c['file_id']}" if c.get("file_id") else "")
                        + f"):\n{c.get('chunk', '')}\n"
                        + (
                            f"Hybrid score: {c['hybrid']:.2f}"
                            if c.get("hybrid") is not None
                            else ""
                        )
                        + (
                            f", Rerank score: {c['rerank']:.2f}"
                            if c.get("rerank") is not None
                            else ""
                        )
                        + "\n\n"
                    )
                    context_idx += 1
        else:
            # Original flat format (no decomposition or single sub-query)
            context_str = "\n\n".join(
                (
                    f"Context {i+1} (Source: {c.get('source', 'unknown')}"
                    + (f", Page: {c['page']}" if c.get("page", 0) > 0 else "")
                    + (f", file_id: {c['file_id']}" if c.get("file_id") else "")
                    + f"):\n{c.get('chunk', '')}\n"
                    + (
                        f"Hybrid score: {c['hybrid']:.2f}"
                        if c.get("hybrid") is not None
                        else ""
                    )
                    + (
                        f", Rerank score: {c['rerank']:.2f}"
                        if c.get("rerank") is not None
                        else ""
                    )
                )
                for i, c in enumerate(ctx)
            )

        # Return formatted contexts with instructions for the LLM
        # Build instructions based on whether decomposition was used
        if has_decomposition:
            decomp_instruction = (
                f"- The original query was decomposed into {len(sub_queries)} sub-queries for better retrieval. "
                f"Contexts are grouped by sub-query above.\n"
                f"- Use information from ALL sub-query groups to fully answer the ORIGINAL query.\n"
                f"- When comparing entities, ensure you include data from each relevant sub-query group.\n"
            )
        else:
            decomp_instruction = ""

        result = (
            f"Found {len(ctx)} relevant document(s):\n\n"
            f"{context_str}\n\n"
            f"Instructions for DOCUMENT CONTEXTS ABOVE (Context 1-{len(ctx)}):\n"
            f"{decomp_instruction}"
            f"- CRITICAL: Every sentence MUST include at least one citation like [1], [2] that refers to the numbered Context items\n"
            f"- Place bracket citations [n] IMMEDIATELY AFTER each sentence\n"
            f"- Example: 'The revenue was $50M [1]. Sales increased by 20% [2].'\n"
            f"- Use ONLY the information from these contexts to answer\n"
            f"- Do NOT include a 'Sources:' line at the end - sources will be added automatically\n"
            f"- If you also have web_search results in other tool responses, answer those naturally WITHOUT citations\n"
            f"- IMPORTANT: The 'file_id' in context headers is for INTERNAL matching only. NEVER show file_id values in your response to users.\n"
        )

        # Prepend external search suggestion if recommended
        if external_msg:
            print(
                "[SEARCH_DOCUMENTS] BUT!!!!!!!!!!!!!! Prepending external search suggestion"
            )
            result = external_msg + result
            print(f"[SEARCH_DOCUMENTS] External search suggestion:\n{result}")

        # Prepend clarification message if quality was poor
        if clarification_msg:
            result = (
                f"[QUALITY WARNING]\n{clarification_msg}\n\n" f"---\n\n" f"{result}"
            )

        return result
    except Exception as e:
        increment_error(MetricsErrorType.RETRIEVAL_FAILED)
        return f"Error during document retrieval: {str(e)}"


# ============================================================================
# TOOL 2: CALCULATOR
# ============================================================================

CALCULATOR_SCHEMA = {
    "type": "function",
    "function": {
        "name": "calculator",
        "description": (
            "Perform mathematical calculations. "
            "Use this for arithmetic operations, percentages, or any numerical computation. "
            "Supports basic operations (+, -, *, /), percentages, and more."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "expression": {
                    "type": "string",
                    "description": (
                        "Mathematical expression to evaluate. "
                        "Examples: '2+2', '15% of 100', '(100-80)/80 * 100', '1000 * 1.15'"
                    ),
                },
            },
            "required": ["expression"],
        },
    },
}


@traceable
def execute_calculator(args: Dict[str, Any], context: Dict[str, Any]) -> str:
    """
    Execute the calculator tool.

    Args:
        args: Arguments from LLM containing:
            - expression (str): Mathematical expression to evaluate

        context: System context (not used for calculator)

    Returns:
        String result of the calculation
    """
    expression = args.get("expression", "").strip()

    if not expression:
        return "Error: No expression provided"

    try:
        # Handle percentage expressions
        # "15% of 100" -> "0.15 * 100"
        # "20%" -> "0.20"
        processed = _handle_percentage(expression)

        # Safely evaluate the expression
        result = _safe_eval(processed)

        # Format result nicely
        if isinstance(result, float):
            # Remove unnecessary decimals for whole numbers
            if result.is_integer():
                return str(int(result))
            else:
                return f"{result:.6f}".rstrip("0").rstrip(".")
        return str(result)

    except ZeroDivisionError:
        return "Error: Division by zero"
    except Exception as e:
        return f"Error: Invalid mathematical expression - {str(e)}"


def _handle_percentage(expr: str) -> str:
    """
    Convert percentage expressions to Python math.

    Examples:
        "15% of 100" -> "0.15 * 100"
        "20%" -> "0.20"
        "increase by 10%" -> handled in context
    """
    # Pattern: "X% of Y" -> "X/100 * Y"
    expr = re.sub(
        r"(\d+(?:\.\d+)?)\s*%\s+of\s+(\d+(?:\.\d+)?)",
        r"(\1/100) * \2",
        expr,
        flags=re.IGNORECASE,
    )

    # Pattern: "X%" -> "X/100"
    expr = re.sub(r"(\d+(?:\.\d+)?)\s*%", r"(\1/100)", expr)

    return expr


def _safe_eval(expr: str) -> float:
    """
    Safely evaluate a mathematical expression.

    Security: Only allows numbers and basic math operators.
    """
    # Whitelist: only allow numbers, operators, parentheses, dots, spaces
    if not re.match(r"^[\d\s\+\-\*\/\(\)\.\%]+$", expr):
        raise ValueError("Expression contains invalid characters")

    # Additional safety: limit length
    if len(expr) > 200:
        raise ValueError("Expression too long")

    # Evaluate using Python's eval (safe because we validated input)
    # Note: In production, consider using numexpr or ast-based parser
    result = eval(expr, {"__builtins__": {}}, {})

    return float(result)


# ============================================================================
# TOOL 3: WEB SEARCH
# ============================================================================
WEB_SEARCH_SCHEMA = {
    "type": "function",
    "function": {
        "name": "web_search",
        "description": (
            "Search the web for up-to-date information when not found in internal documents. "
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": (
                        "The search query. Be specific and use keywords from the user's question. "
                        "For example: 'Q1 2024 revenue', 'employee vacation policy', 'sales targets'"
                    ),
                },
                "max_results": {
                    "type": "integer",
                    "description": (
                        "Number of web search results to return. "
                        "Use 3-5 for quick lookups, 8-10 for comprehensive answers. Default is 5."
                    ),
                    "default": 5,
                },
            },
            "required": ["query"],
        },
    },
}

# ============================================================================
# TOOL 4: DOWNLOAD FILE
# ============================================================================
DOWNLOAD_FILE_SCHEMA = {
    "type": "function",
    "function": {
        "name": "download_file",
        "description": (
            "Download files or web pages from URLs. "
            "Use this tool to download:\n"
            "- Direct file links (PDF, Excel, images, etc.)\n"
            "- Web page URLs from web_search results (will be saved as HTML files)\n"
            "- Any HTTP/HTTPS URL that user wants to save locally\n"
            "- Files from sites requiring login (with browser automation fallback)\n"
            "After downloading, files can be attached to emails using send_email tool.\n"
            "If direct download fails, browser automation will attempt to navigate and download."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "file_urls": {
                    "type": "array",
                    "items": {
                        "type": "string",
                        "description": "URL to download (can be file URL or web page URL)",
                    },
                    "description": (
                        "List of URLs to download. Supports:\n"
                        "- External file URLs (https://example.com/report.pdf)\n"
                        "- Web page URLs from search results (https://example.com/article)\n"
                        "- Internal document links\n"
                        "- Login-protected pages (browser automation will handle login)"
                    ),
                },
                "task_description": {
                    "type": "string",
                    "description": (
                        "Optional: Natural language description of what to download. "
                        "Useful for complex downloads like 'December 2025 internet invoice from netbank'. "
                        "This helps the browser automation agent navigate if direct download fails."
                    ),
                },
            },
            "required": ["file_urls"],
        },
    },
}

# ============================================================================
# TOOL 5: SEND EMAIL
# ============================================================================
SEND_EMAIL_SCHEMA = {
    "type": "function",
    "function": {
        "name": "send_email",
        "description": "Send an email with optional attachments.",
        "parameters": {
            "type": "object",
            "properties": {
                "to": {
                    "type": "array",
                    "items": {"type": "string", "format": "email"},
                    "description": "Recipient email addresses",
                },
                "subject": {"type": "string", "description": "Email subject"},
                "body": {"type": "string", "description": "Email body content"},
                "attachments": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": (
                        "Optional file references to attach. ALWAYS use file_id (e.g., 'cmjg8yrab0000xspw5ip79qcb') from tool responses or Available Files list. "
                        "DO NOT use download URLs like '/api/downloads/...' or '/api/files/...'. "
                        "Accepted formats: file_id (recommended), category:filename, or just filename"
                    ),
                    "minItems": 0,
                },
            },
            "required": ["to", "subject", "body"],
        },
    },
}

# ============================================================================
# TOOL 6: CREATE DOCUMENTS
# ============================================================================

CREATE_DOCUMENTS_SCHEMA = {
    "type": "function",
    "function": {
        "name": "create_documents",
        "description": (
            "Format and save text content as downloadable documents (PDF, DOCX, TXT, CSV, XLSX, HTML, MD). "
            "Use this for: saving reports, exporting summaries, creating formatted documents from existing text. "
            "NOTE: This tool only FORMATS content - it does NOT compute or analyze data. "
            "For calculations or data analysis, use code_execution instead."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "documents": {
                    "type": "array",
                    "description": "List of documents to create",
                    "items": {
                        "type": "object",
                        "properties": {
                            "content": {
                                "type": "string",
                                "description": (
                                    "The main content to include in the document. "
                                    "Can include markdown formatting (headings, lists, bold, etc.)"
                                ),
                            },
                            "filename": {
                                "type": "string",
                                "description": (
                                    "Optional filename for the document (without extension). "
                                    "If not provided, title will be used."
                                ),
                            },
                            "title": {
                                "type": "string",
                                "description": "Document title (will be used as filename and header)",
                            },
                            "format": {
                                "type": "string",
                                "enum": [
                                    "pdf",
                                    "docx",
                                    "txt",
                                    "md",
                                    "html",
                                    "csv",
                                    "xlsx",
                                ],
                                "description": (
                                    "Output format:\n"
                                    "- 'pdf': PDF document (formatted, professional)\n"
                                    "- 'docx': Microsoft Word document\n"
                                    "- 'txt': Plain text file\n"
                                    "- 'md': Markdown file (preserves markdown formatting)\n"
                                    "- 'html': HTML file (web-ready)\n"
                                    "- 'csv': CSV file (for tabular data)\n"
                                    "- 'xlsx': Excel spreadsheet (for tabular data)"
                                ),
                                "default": "pdf",
                            },
                            "metadata": {
                                "type": "object",
                                "description": "Optional metadata (author, date, description, etc.)",
                                "properties": {
                                    "author": {"type": "string"},
                                    "subject": {"type": "string"},
                                    "keywords": {"type": "string"},
                                },
                            },
                        },
                        "required": ["content", "title"],
                    },
                    "minItems": 1,
                },
            },
            "required": ["documents"],
        },
    },
}


@traceable
def execute_web_search(args: Dict[str, Any], context: Dict[str, Any]) -> str:
    """
    Execute the web_search tool.

    Args:
        args: Arguments from LLM containing:
            - query (str): Search query
            - top_k (int, optional): Number of results to return
    """
    query = args.get("query", "")
    max_results = args.get("max_results", 5) or Config.WEB_SEARCH_MAX_RESULTS
    try:
        logger.debug(f"[WEB_SEARCH] Using provider: {Config.WEB_SEARCH_PROVIDER}")
        web_search_service = WebSearchService(
            provider=Config.WEB_SEARCH_PROVIDER,
            max_results=max_results,
            tavily_api_key=Config.TAVILY_API_KEY,
        )
        web_search_results = web_search_service.search(
            query=query, max_results=max_results
        )
        if web_search_results:
            return web_search_service.format_for_agent(web_search_results, query)
        else:
            print("[WEB_SEARCH] No results found, falling back to document contexts")
            return "No relevant web results found."
    except Exception as e:
        increment_error(MetricsErrorType.WEB_SEARCH_FAILED)
        return f"Error during web search: {str(e)}"


async def _try_browser_download(
    url: str,
    user_id: str,
    context: Dict[str, Any],
    task_description: str = "",
) -> str | None:
    """
    Try to download a file using browser automation as fallback.

    Args:
        url: The URL that failed with HTTP download
        user_id: User ID for directory organization
        context: System context with file_service
        task_description: Optional natural language description of what to download

    Returns:
        Success message string if download succeeded, None if failed
    """
    if not Config.BROWSER_USE_ENABLED:
        logger.debug(
            "[BROWSER_DOWNLOAD] Browser automation disabled, skipping fallback"
        )
        return None

    try:
        if not is_browser_use_available():
            logger.warning("[BROWSER_DOWNLOAD] browser-use not available")
            return None

        logger.info(f"[BROWSER_DOWNLOAD] Attempting browser fallback for: {url}")

        # Build task description for the browser agent
        # Note: Don't add URL here - browser_download's _build_task_prompt handles it
        if task_description:
            task = task_description
        else:
            task = "Download the file from this URL"

        logger.debug(f"[BROWSER_DOWNLOAD] Task description: {task}")

        # Download directory
        download_dir = os.path.join(Config.DOWNLOAD_BASE, user_id)

        # Get credentials if configured
        credentials = None
        if Config.BROWSER_TEST_USERNAME and Config.BROWSER_TEST_PASSWORD:
            credentials = {
                "username": Config.BROWSER_TEST_USERNAME,
                "password": Config.BROWSER_TEST_PASSWORD,
            }

        # Run browser download
        success, message, file_path = await browser_download(
            task=task,
            download_dir=download_dir,
            url=url,
            credentials=credentials,
        )

        if not success or not file_path:
            logger.warning(f"[BROWSER_DOWNLOAD] Failed: {message}")
            return None

        # Register the downloaded file
        file_service: FileService = context.get("file_service")
        if not file_service:
            logger.error("[BROWSER_DOWNLOAD] FileService not in context")
            return None

        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)

        # Guess mime type from extension
        import mimetypes

        mime_type, _ = mimetypes.guess_type(file_path)
        mime_type = mime_type or "application/octet-stream"

        result = await file_service.register_file(
            user_email=user_id,
            category="downloaded",
            original_name=filename,
            storage_path=file_path,
            source_tool="download_file_browser",
            mime_type=mime_type,
            size_bytes=file_size,
            source_url=url,
            conversation_id=context.get("conversation_id"),
            dept_id=context.get("dept_id"),
            metadata={"file_for_user": True, "browser_download": True},
        )

        file_id = result["file_id"]
        download_url = result["download_url"]

        logger.info(f"[BROWSER_DOWNLOAD] Success: {filename} -> {file_id}")

        return (
            f"✅ Downloaded (via browser): {filename}\n"
            f"   File ID: {file_id}\n"
            f"   Download: [{filename}]({download_url})\n"
            f"   → Use file ID '{file_id}' for email attachments"
        )

    except Exception as e:
        logger.error(f"[BROWSER_DOWNLOAD] Error: {e}")
        return None


@traceable
async def execute_download_file(args: Dict[str, Any], context: Dict[str, Any]) -> str:
    """
    Execute the download_file tool - Downloads files from URLs to server and returns download links.

    Architecture:
    =============
    User asks "download this PDF" → LLM calls download_file → Server downloads →
    Server saves to downloads/{user_id}/{timestamp}_{filename} → Returns download URL →
    User sees clickable link in chat

    Args:
        args: Arguments from LLM containing:
            - file_urls (list[str]): List of URLs to download

        context: System context containing:
            - user_id (str): Current user ID for directory organization
            - dept_id (str): Department ID (optional, for multi-tenancy)
            - file_service (FileService): Injected file service for registration

    Returns:
        str: Success message with download links OR error message

    Example output:
        "Downloaded 2 files:
        👉 [南京著名景点简介.pdf](/api/files/clx_abc123)
        👉 [销售报告.xlsx](/api/files/clx_def456)"

    TODO Implementation Steps:
    ==========================
    Step 1: Validate inputs
        - Check if file_urls list is not empty
        - Get user_id from context (required for directory structure)
        - Return error if user_id missing: "Error: User ID not found in context"

    Step 2: Setup downloads directory
        - Create base downloads directory: DOWNLOAD_BASE = "downloads"
        - Create user-specific subdirectory: downloads/{user_id}/
        - Use os.makedirs(path, exist_ok=True) to ensure directory exists
        - Example: downloads/user123/

    Step 3: Download each file with safety checks
        For each file_url in file_urls:
            a) Validate URL format
                - Use urllib.parse.urlparse to check scheme is http/https
                - Reject file:// or other dangerous schemes
                - Return error: "Invalid URL scheme: {url}"

            b) Check file size before downloading (HEAD request)
                - Make HEAD request: requests.head(url, timeout=10)
                - Get Content-Length header
                - MAX_DOWNLOAD_SIZE = 100 * 1024 * 1024  # 100MB
                - If size > MAX_DOWNLOAD_SIZE: skip with error message

            c) Download file content
                - Use requests.get(url, stream=True, timeout=30)
                - Check response.status_code == 200
                - IMPORTANT: Must use stream=True to enable chunked reading

            d) Extract filename and sanitize
                - Try Content-Disposition header first
                - Fallback to URL path: url.split("/")[-1]
                - Decode URL encoding: urllib.parse.unquote(filename)
                - Sanitize: Remove dangerous characters, path traversal attempts
                - Replace spaces/special chars with underscores

            e) Generate unique filename
                - Add timestamp prefix: datetime.now().strftime("%Y%m%d_%H%M%S")
                - Format: {timestamp}_{sanitized_filename}
                - Example: "20250118_120530_report.pdf"

            f) Save file to disk
                - Full path: downloads/{user_id}/{timestamped_filename}
                - MUST use chunked writing to preserve file size and save memory:
                  with open(path, "wb") as f:
                      for chunk in response.iter_content(chunk_size=8192):
                          if chunk:
                              f.write(chunk)
                - DO NOT use: f.write(response.content) - loads entire file into memory

            g) Register file in FileRegistry database
                - Call file_service.register_file() to get file_id and download_url
                - Format: /api/files/{file_id} (unified for all file types)
                - Store in results list: (original_url, download_url, filename, file_id)

    Step 4: Build response message for LLM
        - Count successful downloads
        - Format markdown links: [filename](download_url)
        - Add emoji: "👉 [filename.pdf](url)"
        - Handle errors: "⚠️ Failed to download: {url} - {error}"
        - Return formatted string that LLM will show to user

    Step 5: Error handling
        - Wrap entire function in try/except
        - Catch requests.RequestException for network errors
        - Catch IOError for file system errors
        - Return user-friendly error messages
        - Log errors with logger.error() for debugging

    Security Considerations:
    ========================
    - Validate URL schemes (only http/https)
    - Sanitize filenames (no path traversal: ../, \\)
    - Check file size limits (prevent disk fill attacks)
    - Use user-specific directories (prevent file access across users)
    - Set download timeout (prevent hanging requests)
    - Validate content types if needed (optional)

    Directory Structure:
    ====================
    downloads/
    ├── user123/
    │   ├── 20250118_120530_report.pdf
    │   ├── 20250118_120545_image.png
    │   └── ...
    └── user456/
        └── ...

    Example Usage:
    ==============
    User: "Please download this PDF: https://example.com/report.pdf"
    LLM calls: download_file(file_urls=["https://example.com/report.pdf"])
    Server downloads → Saves to downloads/user123/20250118_120530_report.pdf
    Server registers in FileRegistry → file_id: clx_abc123, download_url: /api/files/clx_abc123
    Returns: "Downloaded 1 file:\\n👉 [report.pdf](/api/files/clx_abc123) [file:clx_abc123]"
    LLM shows: "Downloaded successfully! 👉 [report.pdf](/api/files/clx_abc123)"
    """
    # For now, return placeholder
    file_list = args.get("file_urls", [])
    if not file_list:
        return "No file URLs provided"

    user_id = context.get("user_id")
    if not user_id:
        return "Error: User ID not found in context"

    os.makedirs(f"{Config.DOWNLOAD_BASE}/{user_id}/", exist_ok=True)
    results = []
    errors = []

    # Use browser-like headers to avoid 403 errors from websites
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
        "Accept-Encoding": "gzip, deflate, br",
        "Connection": "keep-alive",
    }

    for file_url in file_list:
        try:
            logger.debug(f"[DOWNLOAD_FILE] File URL: {file_url}")
            parsed_url = urlparse(file_url)
            if parsed_url.scheme not in ("http", "https"):
                errors.append(f"⚠️ Invalid URL scheme: {file_url}")
                continue

            # Check file size before downloading
            try:
                request_header = requests.head(
                    file_url, timeout=10, allow_redirects=True, headers=headers
                )
                content_length = request_header.headers.get("Content-Length", 0)
                if Config.MAX_DOWNLOAD_SIZE_MB and content_length:
                    file_size_mb = int(content_length) / (1024 * 1024)
                    if file_size_mb > Config.MAX_DOWNLOAD_SIZE_MB:
                        errors.append(
                            f"⚠️ File too large: {file_url} ({file_size_mb:.1f}MB > {Config.MAX_DOWNLOAD_SIZE_MB}MB)"
                        )
                        continue
            except requests.RequestException:
                # HEAD request failed, try downloading anyway
                pass

            # Download file with browser-like headers
            res = requests.get(
                file_url,
                stream=True,
                timeout=Config.DOWNLOAD_TIMEOUT,
                allow_redirects=True,
                headers=headers,
            )
            if res.status_code != 200:
                # Try browser fallback for auth errors (401, 403) or other failures
                if res.status_code in (401, 403, 404, 500, 502, 503):
                    browser_result = await _try_browser_download(
                        file_url, user_id, context, args.get("task_description", "")
                    )
                    if browser_result:
                        results.append(browser_result)
                        continue
                errors.append(f"⚠️ Download failed: {file_url} (HTTP {res.status_code})")
                continue

            # Extract filename
            filename = ""
            if "Content-Disposition" in res.headers:
                cd = res.headers.get("Content-Disposition")
                fname_match = re.findall('filename="?([^\'";]+)"?', cd)
                if fname_match:
                    filename = fname_match[0]

            if not filename:
                filename = os.path.basename(parsed_url.path)

            if not filename:
                filename = "downloaded_file"

            filename = urllib.parse.unquote(filename)
            filename = re.sub(r"[^\w\-_\. ]", "_", filename)

            # Add extension based on Content-Type if filename has no extension
            if "." not in filename or filename.endswith("_"):
                content_type = (
                    res.headers.get("Content-Type", "").split(";")[0].strip().lower()
                )
                extension_map = {
                    "text/html": ".html",
                    "application/pdf": ".pdf",
                    "image/jpeg": ".jpg",
                    "image/png": ".png",
                    "image/gif": ".gif",
                    "image/webp": ".webp",
                    "application/json": ".json",
                    "text/plain": ".txt",
                    "text/csv": ".csv",
                    "application/xml": ".xml",
                    "text/xml": ".xml",
                    "application/zip": ".zip",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
                    "application/msword": ".doc",
                    "application/vnd.ms-excel": ".xls",
                }
                ext = extension_map.get(content_type, "")
                if ext:
                    filename = filename.rstrip("_") + ext
                    logger.info(
                        f"[DOWNLOAD_FILE] Added extension {ext} based on Content-Type: {content_type}"
                    )

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_filename = f"{timestamp}_{filename}"
            download_path = os.path.join(Config.DOWNLOAD_BASE, user_id, unique_filename)

            # Save file in chunks
            with open(download_path, "wb") as f:
                for chunk in res.iter_content(chunk_size=8192):
                    if chunk:  # Filter out keep-alive chunks
                        f.write(chunk)

            # Register file in FileRegistry database using FileService
            file_size = os.path.getsize(download_path)
            mime_type = res.headers.get("Content-Type", "application/octet-stream")

            file_service: FileService = context.get("file_service")
            if not file_service:
                raise ValueError("FileService not provided in context")

            result = await file_service.register_file(
                user_email=user_id,
                category="downloaded",
                original_name=filename,
                storage_path=download_path,
                source_tool="download_file",
                mime_type=mime_type,
                size_bytes=file_size,
                source_url=file_url,
                conversation_id=context.get("conversation_id"),
                dept_id=context.get("dept_id"),
                metadata={"file_for_user": True},
            )

            file_id = result["file_id"]
            download_url = result["download_url"]
            # Put file_id FIRST and most prominent for attachment operations
            results.append(
                f"✅ Downloaded: {filename}\n"
                f"   File ID: {file_id}\n"
                f"   Download: [{filename}]({download_url})\n"
                f"   → Use file ID '{file_id}' for email attachments"
            )

        except requests.RequestException as e:
            # Try browser fallback if enabled
            browser_result = await _try_browser_download(
                file_url, user_id, context, args.get("task_description", "")
            )
            if browser_result:
                results.append(browser_result)
            else:
                errors.append(f"⚠️ Network error: {file_url} - {str(e)}")
                logger.error(f"[DOWNLOAD_FILE] Request error for {file_url}: {e}")
        except IOError as e:
            errors.append(f"⚠️ File write error: {file_url} - {str(e)}")
            logger.error(f"[DOWNLOAD_FILE] IO error for {file_url}: {e}")
        except Exception as e:
            errors.append(f"⚠️ Unexpected error: {file_url} - {str(e)}")
            logger.error(f"[DOWNLOAD_FILE] Unexpected error for {file_url}: {e}")

    # Build response message
    if not results and not errors:
        return "No files were downloaded"

    response_parts = []
    if results:
        response_parts.append(
            f"Downloaded {len(results)} file(s):\n" + "\n".join(results)
        )
    if errors:
        response_parts.append("\nErrors:\n" + "\n".join(errors))

    return "\n".join(response_parts)


@traceable
async def execute_send_email(args: Dict[str, Any], context: Dict[str, Any]) -> str:
    """
    Execute the send_email tool with unified file resolution.

    Args:
        args: Arguments from LLM containing:
            - to (list[str]): Recipient email addresses
            - subject (str): Email subject
            - body (str): Email body content
            - attachments (list[str], optional): Attachment identifiers
              Supports multiple formats:
                - "file_abc123" or "clx_abc123": File ID from FileRegistry (recommended)
                - "chat:report.pdf": Category:filename format
                - "report.pdf": Filename only (searches user's files)

        context: System context containing:
            - user_id (str): Current user email
            - dept_id (str): Department ID for shared file access
            - file_service (FileService): Injected file service

    Returns:
        Success/error message string
    """
    to_addresses = args.get("to", [])
    subject = args.get("subject", "")
    body = args.get("body", "")
    if not to_addresses or not subject or not body:
        return "Missing required email fields"

    user_email = context.get("user_id")
    dept_id = context.get("dept_id")
    if not user_email:
        return "Error: User email not found in context"

    file_service: FileService = context.get("file_service")
    if not file_service:
        return "Error: FileService not provided in context"

    # Get attachment references from LLM args
    attachment_refs = args.get("attachments", [])
    logger.debug(f"[SEND_EMAIL] Attachment references: {attachment_refs}")

    real_attachments = []
    errors = []

    # Resolve each attachment reference using FileService
    for ref in attachment_refs:
        try:
            file_path = await file_service.get_file_path(
                ref, user_email, dept_id=dept_id
            )
            real_attachments.append(file_path)
        except FileNotFoundError:
            errors.append(f"⚠️ File not found: {ref}")
            logger.error(f"[SEND_EMAIL] File not found: {ref}")
        except Exception as e:
            errors.append(f"⚠️ Error resolving {ref}: {str(e)}")
            logger.error(f"[SEND_EMAIL] Unexpected error resolving {ref}: {e}")

    if errors:
        return "Failed to send email:\n" + "\n".join(errors)

    logger.debug(f"[SEND_EMAIL] Resolved attachments: {real_attachments}")

    # Send email with resolved file paths
    result = send_email(to_addresses, subject, body, real_attachments)

    if result.get("status", "failed") == "failed":
        logger.error(
            f"[SEND_EMAIL] Failed to send email: {result.get('error', 'Unknown error')}"
        )
        return f"Failed to send email: {result.get('error', 'Unknown error')}"

    return f"Email sent successfully to recipients: {', '.join(to_addresses)} with subject: {subject}"


@traceable
async def execute_create_documents(
    args: Dict[str, Any], context: Dict[str, Any]
) -> str:
    """
    Execute the create_documents tool - supports creating multiple documents in one call.

    Args:
        args: Arguments from LLM containing:
            - documents (list): List of document objects, each containing:
                - content (str): Main document content (supports markdown)
                - title (str): Document title
                - format (str): Output format - pdf/docx/txt/md/html/csv/xlsx (default: 'pdf')
                - metadata (dict, optional): Author, subject, keywords, etc.

        context: System context containing:
            - user_id (str): User ID for file ownership
            - file_service (FileService): Injected file service

    Returns:
        Formatted string with download links for all created documents

    Implementation TODO:

    SUPPORTED FORMATS (7 total):

    1. PDF - reportlab (pip install reportlab)
       - Professional formatted documents
       - Markdown → styled PDF with fonts/sizes
       - Page numbers, headers, metadata

    2. DOCX - python-docx (pip install python-docx)
       - Microsoft Word documents
       - Markdown → Word styles (headings, lists, bold)
       - Document properties

    3. TXT - built-in
       - Plain text, strip markdown or keep it
       - open(path, "w", encoding="utf-8")

    4. MD - built-in
       - Markdown file (save content as-is)
       - open(path, "w", encoding="utf-8")

    5. HTML - markdown2 (pip install markdown2)
       - Convert markdown → HTML
       - Wrap in proper HTML structure
       - Add CSS styling

    6. CSV - built-in csv module
       - For tabular data
       - Parse content, write with csv.writer
       - Headers, rows

    7. XLSX - openpyxl (pip install openpyxl)
       - Excel spreadsheets
       - Parse content for tables
       - Format cells (bold headers, borders)

    IMPLEMENTATION:
    Loop through documents → validate → generate based on format → collect results

    LIBRARIES:
    pip install reportlab python-docx markdown2 openpyxl
    """
    documents = args.get("documents", [])
    if not documents:
        return "No documents provided for creation"

    user_id = context.get("user_id")
    if not user_id:
        return "Error: User ID not found in context"

    file_service: FileService = context.get("file_service")
    if not file_service:
        return "Error: FileService not provided in context"

    # Create user downloads directory
    user_dir = os.path.join(Config.DOWNLOAD_BASE, user_id)
    os.makedirs(user_dir, exist_ok=True)

    results = []
    errors = []

    for doc in documents:
        # Skip None or invalid document entries
        if doc is None or not isinstance(doc, dict):
            errors.append("⚠️ Skipped invalid document entry (None or non-dict)")
            continue

        try:
            content = doc.get("content", "")
            title = doc.get("title", "untitled")
            custom_filename = doc.get("filename", "")
            format_type = doc.get("format", "pdf").lower()
            metadata = doc.get("metadata", {})

            # Validate required fields
            if not content or not content.strip():
                errors.append(f"⚠️ Skipped document '{title}': no content provided")
                continue

            if format_type not in ("pdf", "docx", "txt", "md", "html", "csv", "xlsx"):
                errors.append(f"⚠️ Unsupported format: {format_type} for {title}")
                continue

            # Sanitize filename
            base_filename = custom_filename if custom_filename else title
            base_filename = re.sub(r"[^\w\-_\. ]", "_", base_filename)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # Generate file based on format
            if format_type == "txt":
                filename = f"{timestamp}_{base_filename}.txt"
                filepath = os.path.join(user_dir, filename)
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(content)

            elif format_type == "md":
                filename = f"{timestamp}_{base_filename}.md"
                filepath = os.path.join(user_dir, filename)
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(content)

            elif format_type == "html":
                filename = f"{timestamp}_{base_filename}.html"
                filepath = os.path.join(user_dir, filename)
                try:
                    import markdown2

                    html_content = markdown2.markdown(content)
                    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
        h1, h2, h3 {{ color: #333; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    {html_content}
</body>
</html>"""
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(full_html)
                except ImportError:
                    errors.append(f"⚠️ HTML format requires markdown2 library: {title}")
                    continue

            elif format_type == "pdf":
                filename = f"{timestamp}_{base_filename}.pdf"
                filepath = os.path.join(user_dir, filename)
                try:
                    from fpdf import FPDF

                    # Create PDF with Unicode support
                    pdf = FPDF(orientation="P", unit="mm", format="A4")
                    pdf.set_margins(left=20, top=20, right=20)
                    pdf.set_auto_page_break(auto=True, margin=20)

                    # Try to add Unicode font for CJK support
                    unicode_font = None
                    font_paths = [
                        # Windows - use .ttf not .ttc for better compatibility
                        ("C:/Windows/Fonts/simhei.ttf", "SimHei"),
                        ("C:/Windows/Fonts/msyh.ttc", "MSYH"),
                        ("C:/Windows/Fonts/simsun.ttc", "SimSun"),
                        # Linux
                        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "DejaVu"),
                        # macOS
                        ("/Library/Fonts/Arial Unicode.ttf", "ArialUnicode"),
                    ]
                    for font_path, font_name in font_paths:
                        if os.path.exists(font_path):
                            try:
                                pdf.add_font(font_name, fname=font_path)
                                unicode_font = font_name
                                logger.info(f"[PDF] Loaded font: {font_name}")
                                break
                            except Exception as fe:
                                logger.warning(f"[PDF] Font {font_name} failed: {fe}")
                                continue

                    pdf.add_page()

                    # Calculate usable width
                    usable_width = pdf.w - pdf.l_margin - pdf.r_margin

                    # Use unicode font or fallback
                    def set_font_safe(size, style=""):
                        if unicode_font:
                            pdf.set_font(unicode_font, size=size)
                        else:
                            pdf.set_font("Helvetica", style=style, size=size)

                    # Add title
                    set_font_safe(16, "B")
                    pdf.set_text_color(33, 33, 33)
                    pdf.multi_cell(w=usable_width, h=10, text=title, align="L")
                    pdf.ln(5)

                    # Add separator line
                    pdf.set_draw_color(150, 150, 150)
                    y = pdf.get_y()
                    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
                    pdf.ln(8)

                    # Add content
                    set_font_safe(10)
                    pdf.set_text_color(51, 51, 51)

                    # Helper function to render a markdown table
                    def render_table(table_lines):
                        if len(table_lines) < 2:
                            return

                        # Parse table rows
                        rows = []
                        for tl in table_lines:
                            # Skip separator rows (|---|---|)
                            if tl.replace("|", "").replace("-", "").replace(":", "").strip() == "":
                                continue
                            # Parse cells
                            cells = [c.strip() for c in tl.split("|")]
                            # Remove empty first/last cells from leading/trailing |
                            if cells and cells[0] == "":
                                cells = cells[1:]
                            if cells and cells[-1] == "":
                                cells = cells[:-1]
                            if cells:
                                rows.append(cells)

                        if not rows:
                            return

                        # Calculate column widths
                        num_cols = max(len(r) for r in rows)
                        col_width = usable_width / num_cols if num_cols > 0 else usable_width
                        row_height = 8

                        # Draw table
                        pdf.set_draw_color(100, 100, 100)
                        for row_idx, row in enumerate(rows):
                            # Pad row to have consistent columns
                            while len(row) < num_cols:
                                row.append("")

                            # Header row (first row) - bold with background
                            if row_idx == 0:
                                set_font_safe(10, "B")
                                pdf.set_fill_color(240, 240, 240)
                                for cell in row:
                                    pdf.cell(w=col_width, h=row_height, text=cell, border=1, fill=True, align="C")
                                pdf.ln(row_height)
                                set_font_safe(10)
                            else:
                                # Data rows
                                for cell in row:
                                    pdf.cell(w=col_width, h=row_height, text=cell, border=1, align="C")
                                pdf.ln(row_height)

                        pdf.ln(4)  # Space after table

                    # Process content line by line, collecting table lines
                    lines = content.split("\n")
                    i = 0
                    while i < len(lines):
                        line = lines[i].strip()

                        if not line:
                            pdf.ln(4)
                            i += 1
                            continue

                        # Check if this is a markdown table row
                        if line.startswith("|") and line.endswith("|"):
                            # Collect all consecutive table lines
                            table_lines = []
                            while i < len(lines):
                                tl = lines[i].strip()
                                if tl.startswith("|") and (tl.endswith("|") or tl.rstrip().endswith("|")):
                                    table_lines.append(tl)
                                    i += 1
                                else:
                                    break
                            render_table(table_lines)
                            continue

                        # Strip markdown formatting
                        clean = line.replace("**", "").replace("*", "")

                        # Handle markdown headers
                        if line.startswith("### "):
                            set_font_safe(11, "B")
                            pdf.multi_cell(w=usable_width, h=7, text=clean[4:])
                            set_font_safe(10)
                        elif line.startswith("## "):
                            pdf.ln(3)
                            set_font_safe(13, "B")
                            pdf.multi_cell(w=usable_width, h=8, text=clean[3:])
                            set_font_safe(10)
                        elif line.startswith("# "):
                            pdf.ln(4)
                            set_font_safe(14, "B")
                            pdf.multi_cell(w=usable_width, h=9, text=clean[2:])
                            set_font_safe(10)
                        elif line.startswith("- ") or line.startswith("* "):
                            pdf.multi_cell(w=usable_width, h=6, text=f"  • {clean[2:]}")
                        else:
                            pdf.multi_cell(w=usable_width, h=6, text=clean)

                        i += 1

                    pdf.output(filepath)
                except ImportError:
                    errors.append(f"⚠️ PDF format requires fpdf2 library: {title}")
                    continue
                except Exception as pdf_err:
                    logger.error(f"[PDF] Creation failed: {pdf_err}")
                    errors.append(
                        f"⚠️ PDF creation failed for '{title}': {str(pdf_err)}"
                    )
                    continue

            elif format_type == "docx":
                filename = f"{timestamp}_{base_filename}.docx"
                filepath = os.path.join(user_dir, filename)
                try:
                    from docx import Document

                    doc_docx = Document()

                    # Add title
                    doc_docx.add_heading(title, 0)

                    # Add content (basic paragraph formatting)
                    for line in content.split("\n"):
                        if line.strip():
                            doc_docx.add_paragraph(line)

                    # Set metadata if provided
                    core_props = doc_docx.core_properties
                    if metadata.get("author"):
                        core_props.author = metadata["author"]
                    if metadata.get("subject"):
                        core_props.subject = metadata["subject"]

                    doc_docx.save(filepath)
                except ImportError:
                    errors.append(
                        f"⚠️ DOCX format requires python-docx library: {title}"
                    )
                    continue

            elif format_type == "csv":
                filename = f"{timestamp}_{base_filename}.csv"
                filepath = os.path.join(user_dir, filename)
                import csv

                # Simple CSV: treat each line as a row, split by commas or tabs
                with open(filepath, "w", encoding="utf-8", newline="") as f:
                    writer = csv.writer(f)
                    for line in content.split("\n"):
                        if line.strip():
                            # Try tab-separated first, then comma-separated
                            if "\t" in line:
                                cells = [cell.strip() for cell in line.split("\t")]
                            else:
                                cells = [cell.strip() for cell in line.split(",")]
                            writer.writerow(cells)

            elif format_type == "xlsx":
                filename = f"{timestamp}_{base_filename}.xlsx"
                filepath = os.path.join(user_dir, filename)
                try:
                    from openpyxl import Workbook

                    wb = Workbook()
                    ws = wb.active
                    ws.title = title[:31]  # Excel sheet name limit

                    # Parse content as simple table (comma or tab-separated)
                    row_num = 1
                    for line in content.split("\n"):
                        if line.strip():
                            if "\t" in line:
                                cells = [cell.strip() for cell in line.split("\t")]
                            else:
                                cells = [cell.strip() for cell in line.split(",")]
                            for col_num, cell_value in enumerate(cells, 1):
                                ws.cell(row=row_num, column=col_num, value=cell_value)
                            row_num += 1

                    wb.save(filepath)
                except ImportError:
                    errors.append(f"⚠️ XLSX format requires openpyxl library: {title}")
                    continue

            # Register file using FileService
            file_size = os.path.getsize(filepath)
            mime_type_map = {
                "pdf": "application/pdf",
                "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "txt": "text/plain",
                "md": "text/markdown",
                "html": "text/html",
                "csv": "text/csv",
                "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            }

            result = await file_service.register_file(
                user_email=user_id,
                category="created",
                original_name=f"{base_filename}.{format_type}",
                storage_path=filepath,
                source_tool="create_documents",
                mime_type=mime_type_map.get(format_type, "application/octet-stream"),
                size_bytes=file_size,
                conversation_id=context.get("conversation_id"),
                dept_id=context.get("dept_id"),
                metadata={
                    "title": title,
                    "format": format_type,
                    "file_for_user": True,
                },
            )

            file_id = result["file_id"]
            download_url = result["download_url"]
            # Put file_id FIRST and most prominent for attachment operations
            results.append(
                f"✅ Created: {base_filename}.{format_type}\n"
                f"   File ID: {file_id}\n"
                f"   Download: [{base_filename}.{format_type}]({download_url})\n"
                f"   → Use file ID '{file_id}' for email attachments"
            )
            logger.info(
                f"[CREATE_DOCUMENTS] Created {format_type} document: {filepath} | file_id: {file_id}"
            )

        except Exception as e:
            errors.append(
                f"⚠️ Failed to create document '{doc.get('title', 'unknown')}': {str(e)}"
            )
            logger.error(f"[CREATE_DOCUMENTS] Error creating document: {e}")

    # Build response message
    if not results and not errors:
        return "No documents were created"

    response_parts = []
    if results:
        response_parts.append(
            f"Created {len(results)} document(s):\n" + "\n".join(results)
        )
    if errors:
        response_parts.append("\nErrors:\n" + "\n".join(errors))

    return "\n".join(response_parts)


# ============================================================================
# TOOL 7: CODE EXECUTION (E2B Sandbox)
# ============================================================================

CODE_EXECUTION_SCHEMA = {
    "type": "function",
    "function": {
        "name": "code_execution",
        "description": (
            "Execute Python code for computation and data analysis. "
            "Use this when you need to: calculate statistics, analyze data, transform datasets, "
            "process numbers, or run algorithms. Available: pandas, numpy, matplotlib, math. "
            "CRITICAL: Do NOT create/write files (no to_excel, to_csv, open, etc). "
            "Just compute data and return it as the last expression or print(). "
            "For file creation, use create_documents tool instead."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "code": {
                    "type": "string",
                    "description": (
                        "Python code to execute. The last expression value is captured (like Jupyter). "
                        "Available libraries: pandas, numpy, matplotlib, math, json, datetime. "
                        "IMPORTANT: Do NOT use to_excel(), to_csv(), or any file-writing functions. "
                        "Just return computed data (e.g., df.to_dict('records') or print(result))."
                    ),
                },
            },
            "required": ["code"],
        },
    },
}


@traceable
async def execute_code_execution(args: Dict[str, Any], context: Dict[str, Any]) -> str:
    """
    Execute Python code in E2B sandbox.

    Args:
        args: Arguments from LLM containing:
            - code (str): Python code to execute

        context: System context (may contain data from previous steps)

    Returns:
        String result from code execution (stdout + any errors)
    """
    code = args.get("code", "").strip()

    if not code:
        return "Error: No code provided"

    if not Config.CODE_EXECUTION_ENABLED:
        return "Error: Code execution is disabled."

    # Check if E2B is configured
    if not Config.E2B_API_KEY:
        return (
            "Error: E2B API key not configured. Set E2B_API_KEY environment variable."
        )

    try:
        from e2b_code_interpreter import Sandbox

        # Execute code in E2B sandbox
        # Note: E2B v2.x uses Sandbox.create() and reads API key from E2B_API_KEY env var
        with Sandbox.create() as sandbox:
            execution = sandbox.run_code(code)

            # Build result from execution
            result_parts = []

            # Capture stdout
            if execution.text:
                result_parts.append(execution.text)

            # Capture any logs/prints
            if execution.logs and execution.logs.stdout:
                for log in execution.logs.stdout:
                    result_parts.append(log)

            # Capture errors
            if execution.error:
                result_parts.append(f"Error: {execution.error}")

            if execution.logs and execution.logs.stderr:
                for err in execution.logs.stderr:
                    result_parts.append(f"stderr: {err}")

            if not result_parts:
                logger.info("[CODE_EXECUTION] No output captured from code execution")
                return "Code executed successfully (no output)"

            result = "\n".join(result_parts)
            logger.info(f"[CODE_EXECUTION] Execution result: {result[:150]}...")

            return result

    except ImportError:
        return "Error: e2b-code-interpreter package not installed. Run: pip install e2b-code-interpreter"
    except Exception as e:
        logger.error(f"[CODE_EXECUTION] Error: {e}")
        return f"Error executing code: {str(e)}"


# ============================================================================
# TOOL REGISTRY
# ============================================================================

TOOL_REGISTRY = {
    "search_documents": execute_search_documents,
    "web_search": execute_web_search,
    "calculator": execute_calculator,
    "download_file": execute_download_file,
    "send_email": execute_send_email,
    "create_documents": execute_create_documents,
    "code_execution": execute_code_execution,
}
"""
Maps tool names to their execution functions.
When LLM calls a tool, we lookup the function here and execute it.
"""


# ============================================================================
# ALL TOOLS LIST
# ============================================================================

ALL_TOOLS = [
    SEARCH_DOCUMENTS_SCHEMA,
    CALCULATOR_SCHEMA,
    WEB_SEARCH_SCHEMA,
    DOWNLOAD_FILE_SCHEMA,
    SEND_EMAIL_SCHEMA,
    CREATE_DOCUMENTS_SCHEMA,
    CODE_EXECUTION_SCHEMA,
]
"""
List of all tool schemas to send to OpenAI API.
This tells the LLM what tools are available.
"""


# ============================================================================
# LANGGRAPH TOOL LISTS (Single Source of Truth)
# ============================================================================
# LangGraph chat_completion_with_tools expects a list of tool schemas.
# These are the canonical definitions - import these in langgraph_nodes.py.

TOOL_CALCULATOR = [CALCULATOR_SCHEMA]
TOOL_WEB_SEARCH = [WEB_SEARCH_SCHEMA]
TOOL_DOWNLOAD_FILE = [DOWNLOAD_FILE_SCHEMA]
TOOL_CREATE_DOCUMENTS = [CREATE_DOCUMENTS_SCHEMA]
TOOL_SEND_EMAIL = [SEND_EMAIL_SCHEMA]
TOOL_CODE_EXECUTION = [CODE_EXECUTION_SCHEMA]


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================


def get_tool_executor(tool_name: str):
    """
    Get the execution function for a tool by name.

    Args:
        tool_name: Name of the tool (e.g., "search_documents")

    Returns:
        Execution function or None if tool not found

    Example:
        executor = get_tool_executor("calculator")
        result = executor({"expression": "2+2"}, {})
    """
    return TOOL_REGISTRY.get(tool_name)


@traceable
async def execute_tool_call(
    tool_name: str, tool_args: Dict[str, Any], context: Dict[str, Any]
) -> str:
    """
    Execute a tool call from the LLM.

    Args:
        tool_name: Name of the tool to execute
        tool_args: Arguments from LLM (parsed from JSON)
        context: System context (vector_db, auth, file_service, etc.)

    Returns:
        String result from tool execution

    Raises:
        ValueError: If tool doesn't exist
    """
    # Look up tool executor
    executor = get_tool_executor(tool_name)

    if not executor:
        raise ValueError(
            f"Tool '{tool_name}' not found in registry. Available tools: {list(TOOL_REGISTRY.keys())}"
        )

    # Ensure args is a dict
    executor_args = tool_args if tool_args else {}

    try:
        # Execute the tool (await if async, call if sync)
        if inspect.iscoroutinefunction(executor):
            result = await executor(executor_args, context)
        else:
            result = executor(executor_args, context)
        return result
    except Exception as e:
        # Log error and return error message
        error_msg = f"Error executing tool '{tool_name}': {str(e)}"
        print(error_msg)  # For debugging
        return error_msg
